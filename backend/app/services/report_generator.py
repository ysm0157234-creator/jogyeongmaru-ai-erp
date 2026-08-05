from __future__ import annotations

import io
from datetime import date
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


COMPANY = {
    "representative": "황수영",
    "birth_date": "1985. 5. 15.",
    "address": "경기도 평택시 진위면 서촌로 38-9",
    "company_name": "농업회사법인 주식회사 조경마루",
    "phone": "010-9377-3058",
    "seed_business_number": "제10-평택-2023-30-01호",
}


def _set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Mm(13)
        section.bottom_margin = Mm(13)
        section.left_margin = Mm(15)
        section.right_margin = Mm(15)


def _set_font(run, size: float = 10.0, bold: bool = False) -> None:
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.bold = bold


def _shade(cell, fill: str = "EDEDED") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_text(cell, text: str, *, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(str(text or ""))
    _set_font(r, 9.5, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_info_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Mm(42)
        cells[1].width = Mm(138)
        _shade(cells[0])
        _set_cell_text(cells[0], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(cells[1], value)
    doc.add_paragraph()


def _add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    _set_font(r, 13, True)


def _add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(str(text or ""))
    _set_font(r, 10.5)


def _add_picture(doc: Document, label: str, data: bytes) -> None:
    p = doc.add_paragraph()
    r = p.add_run(label)
    _set_font(r, 10, True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(io.BytesIO(data), width=Mm(160))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _save(doc: Document) -> bytes:
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_compatible_docx(
    *,
    result_data: dict[str, Any],
    shipment: str,
    quarantine_number: str,
    overall_image: bytes,
    closeup_image: bytes,
    warnings: list[str] | None = None,
) -> bytes:
    classification = result_data.get("classification") or {}
    doc = Document()
    _set_margins(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("품종 생산·수입판매 신고서 검토안")
    _set_font(r, 17, True)

    _add_info_table(doc, [
        ("신고 구분", "수입 판매"),
        ("신고인", COMPANY["representative"]),
        ("법인 명칭", COMPANY["company_name"]),
        ("주소", COMPANY["address"]),
        ("전화번호", COMPANY["phone"]),
        ("작물명", str(result_data.get("korean_name") or result_data.get("matched_name") or "")),
        ("품종명", str(result_data.get("cultivar") or result_data.get("matched_name") or "")),
        ("학명", str(result_data.get("scientific_name") or "")),
        ("원산지", str(result_data.get("origin") or "")),
        ("번식방법", str(result_data.get("propagation_method") or "")),
        ("꽃 색상", str(classification.get("flower_color") or "")),
        ("개화기", str(classification.get("flowering_period") or "")),
        ("성숙 초장", str(classification.get("height") or "")),
        ("Shipment", shipment or "미확인"),
        ("검역합격번호", quarantine_number or "미첨부"),
        ("종자업 등록번호", COMPANY["seed_business_number"]),
    ])

    _add_heading(doc, "품종의 특성 설명")
    _add_body(doc, str(result_data.get("characteristics_draft") or ""))
    _add_heading(doc, "품종육성 과정의 설명")
    _add_body(doc, str(result_data.get("breeding_process_draft") or ""))

    if warnings:
        _add_heading(doc, "검토 및 첨부 상태")
        for item in warnings:
            _add_body(doc, f"• {item}")

    doc.add_page_break()
    _add_heading(doc, "품종의 사진")
    _add_picture(doc, "사진 1. 품종 전체 모습", overall_image)
    _add_picture(doc, "사진 2. 꽃·잎 근접 모습", closeup_image)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"작성일: {date.today().year}년 {date.today().month}월 {date.today().day}일")
    _set_font(r, 9.5)
    return _save(doc)


def build_pdf_summary(
    variety_name: str,
    scientific_name: str,
    shipment: str,
    source_invoice: str,
    quarantine_file: str,
) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    _, height = A4
    y = height - 50
    c.setFont("Helvetica-Bold", 15)
    c.drawString(45, y, "Jogyeongmaru AI ERP - Package Summary")
    y -= 28
    c.setFont("Helvetica", 10)
    for line in [
        f"Variety: {variety_name}",
        f"Scientific name: {scientific_name}",
        f"Shipment: {shipment or 'Not found'}",
        f"Invoice source: {source_invoice or 'Not attached'}",
        f"Quarantine source: {quarantine_file or 'Not attached'}",
        "HWPX and compatible DOCX are included when available.",
    ]:
        c.drawString(45, y, line[:110])
        y -= 16
    c.save()
    return buffer.getvalue()
